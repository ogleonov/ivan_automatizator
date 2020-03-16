import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from Scaler import scale_image
from caption import Figure
from initvars import dirIn
from initvars import dirOut
from initvars import docPath
from initvars import height
from initvars import widht

'To delete the files in dirOut'
folder = dirOut
for filename in os.listdir(folder):
    file_path = os.path.join(folder, filename)
    try:
        if os.path.isfile(file_path) or os.path.islink(file_path):
            os.unlink(file_path)
        elif os.path.isdir(file_path):
            shutil.rmtree(file_path)
    except Exception as e:
        print('Failed to delete %s. Reason: %s' % (file_path, e))


def loadImagesToOutDir(dir):
    if not os.path.exists(dirOut + '\\' + dir):
        os.mkdir(dirOut + '\\' + dir)
    for name in os.listdir(dirIn +'\\' + dir):
        scale_image(os.path.join(dirIn +'\\' + dir, name), os.path.join(dirOut + '\\' + dir, name), widht, height)


if __name__ == '__main__':
    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    listOfDirs = os.listdir(dirIn)
    listOfDirs.sort()
    listOfDirs.reverse()
    print(listOfDirs)
    for tmpDir in listOfDirs:
        loadImagesToOutDir(tmpDir)
        listOfImgs =  os.listdir(dirOut+ '\\'+ tmpDir)
        listOfImgs.sort()
        for img in listOfImgs:
            paragraph_format = document.styles['Normal'].paragraph_format
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing = Pt(0)

            p = document.add_paragraph('')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(os.path.join(dirOut + '\\' + tmpDir, img))

            font.bold = True
            plate = document.add_paragraph('Plate ' + img, style='Normal')
            Figure(plate)
            plate.alignment = WD_ALIGN_PARAGRAPH.CENTER
            d = plate.add_run(' - Lift ' + tmpDir + '- Grid #')

            c = document.add_paragraph()
            cc = c.add_run('[caption]')
            cc.bold = False
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = document.add_paragraph('')

    document.save(docPath)
    os.startfile(docPath)
